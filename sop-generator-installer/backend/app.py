import io
import re
import threading
import webbrowser
import os
import sys
from flask import Flask, render_template_string, request, send_file, make_response, jsonify, session
from docxtpl import DocxTemplate
from waitress import serve

# --- Debug logging to file ---
def debug_log(msg):
    log_path = os.path.join(os.environ.get('LOCALAPPDATA', '.'), 'SOP_Generator', 'debug.log')
    os.makedirs(os.path.dirname(log_path), exist_ok=True)
    with open(log_path, 'a') as f:
        f.write(f"{msg}\n")

debug_log(f"=== App starting ===")
debug_log(f"sys.frozen: {getattr(sys, 'frozen', False)}")
debug_log(f"sys.executable: {sys.executable}")
debug_log(f"os.getcwd(): {os.getcwd()}")


# Import our custom BPMN parser
from bpmn_parser import parse_bpmn_to_sop, extract_metadata_from_bpmn

# Import history manager
from history_manager import HistoryManager

# Import archive manager
from archive_manager import ArchiveManager

# --- Helper function to find bundled files ---
def resource_path(relative_path):
    # Check for Nuitka (sets __compiled__ at module level) or PyInstaller (sets sys.frozen)
    is_nuitka = "__compiled__" in globals()
    is_pyinstaller = getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS')

    if is_pyinstaller:
        base_path = sys._MEIPASS
        debug_log(f"Using PyInstaller _MEIPASS: {base_path}")
    elif is_nuitka or 'Temp' in sys.executable:
        # Nuitka onefile extracts to temp - use executable's directory
        base_path = os.path.dirname(sys.executable)
        debug_log(f"Using Nuitka exe dir: {base_path}")
    else:
        # Running from source
        base_path = os.path.abspath(".")
        debug_log(f"Using source dir: {base_path}")
        # Special handling for template files
        if relative_path in ('final_master_template_2.docx', 'sana_template.docx', 'window_world_template.docx', 'tarabut_template.docx', 'sabah_template.docx'):
            deployment_path = os.path.join(base_path, 'SOP_Generator_Deployment', relative_path)
            if os.path.exists(deployment_path):
                return deployment_path
    result = os.path.join(base_path, relative_path)
    debug_log(f"resource_path('{relative_path}') -> {result} (exists: {os.path.exists(result)})")
    return result

# --- Flask App Initialization ---
app = Flask(__name__)
app.secret_key = 'sop-generator-secret-key-change-in-production'  # For session management

PREVIEW_HTML = r'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>SOP Generator - Preview</title>
    <style>
        body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif; margin: 0; background-color: #f8f9fa; }
        .container { max-width: 900px; margin: 0 auto; padding: 2em; background: white; min-height: 100vh; box-shadow: 0 0 20px rgba(0,0,0,0.08); }
        h1 { color: #343a40; border-bottom: 2px solid #007bff; padding-bottom: 0.5em; margin-top: 0; font-size: 1.4em; }
        h2 { color: #495057; margin-top: 1.5em; font-size: 1.1em; border-bottom: 1px solid #dee2e6; padding-bottom: 0.3em; }
        .form-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 1em; }
        .form-group { display: flex; flex-direction: column; }
        .full-width { grid-column: 1 / -1; }
        label { margin-bottom: 0.3em; font-weight: 600; color: #495057; font-size: 0.9em; }
        input[type="text"], textarea { font-size: 0.95rem; padding: 0.6em; border: 1px solid #ced4da; border-radius: 6px; box-sizing: border-box; }
        input[type="text"]:focus, textarea:focus { border-color: #80bdff; outline: 0; box-shadow: 0 0 0 0.15rem rgba(0,123,255,.25); }
        .auto-filled { background-color: #e8f5e9; }
        .btn-row { display: flex; gap: 1em; margin-top: 2em; justify-content: center; }
        .btn { padding: 0.8em 2em; border: none; border-radius: 8px; font-size: 1em; font-weight: 600; cursor: pointer; }
        .btn-primary { background-color: #007bff; color: white; }
        .btn-primary:hover { background-color: #0056b3; }
        .btn-secondary { background-color: #6c757d; color: white; }
        .btn-secondary:hover { background-color: #545b62; }
        .btn:disabled { opacity: 0.6; cursor: not-allowed; }
        .dynamic-list .entry { display: grid; grid-template-columns: 1fr 2fr auto; gap: 0.5em; margin-bottom: 0.5em; }
        .add-btn { background-color: #28a745; color: white; border: none; border-radius: 6px; padding: 0.4em 1em; cursor: pointer; font-size: 0.85em; }
        .remove-btn { background-color: #dc3545; color: white; border: none; border-radius: 6px; padding: 0.4em 0.8em; cursor: pointer; font-size: 0.85em; }
        .spinner { display: inline-block; border: 3px solid #f3f3f3; border-top: 3px solid #007bff; border-radius: 50%; width: 18px; height: 18px; animation: spin 1s linear infinite; vertical-align: middle; margin-right: 0.5em; }
        @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
        .template-btn { padding: 0.5em 1.2em; border: 2px solid #ced4da; border-radius: 6px; background: white; cursor: pointer; font-weight: 600; font-size: 0.9em; color: #495057; transition: all 0.2s; }
        .template-btn:hover { border-color: #007bff; color: #007bff; }
        .template-btn.active { background: #007bff; color: white; border-color: #007bff; }
    </style>
</head>
<body>
    <div class="container">
        <h1>SOP Generator - Review & Generate</h1>
        <p style="color: #6c757d; margin-top: -0.5em; font-size: 0.9em;">Review the auto-populated metadata below, make any changes, then click Generate.</p>

        <form id="sop-form">

            <div style="display:flex;align-items:center;gap:12px;margin-bottom:20px;padding:14px 16px;background-color:#dee2e6;border-radius:8px;border:2px solid #adb5bd;">
                <span style="margin:0;font-weight:700;color:#212529;font-size:1.1em;">Template:</span>
                <button type="button" id="btn-earthlink" data-template="earthlink" onclick="selectTemplate('earthlink')" style="padding:10px 24px;border:2px solid #007bff;border-radius:8px;background-color:#007bff;color:white;cursor:pointer;font-weight:700;font-size:1.05em;">Earthlink</button>
                <button type="button" id="btn-sana" data-template="sana" onclick="selectTemplate('sana')" style="padding:10px 24px;border:2px solid #6c757d;border-radius:8px;background-color:white;color:#343a40;cursor:pointer;font-weight:700;font-size:1.05em;">SANA</button>
                <button type="button" id="btn-window_world" data-template="window_world" onclick="selectTemplate('window_world')" style="padding:10px 24px;border:2px solid #6c757d;border-radius:8px;background-color:white;color:#343a40;cursor:pointer;font-weight:700;font-size:1.05em;">Window World</button>
                <button type="button" id="btn-tarabut" data-template="tarabut" onclick="selectTemplate('tarabut')" style="padding:10px 24px;border:2px solid #6c757d;border-radius:8px;background-color:white;color:#343a40;cursor:pointer;font-weight:700;font-size:1.05em;">Tarabut</button>
                <button type="button" id="btn-sabah" data-template="sabah" onclick="selectTemplate('sabah')" style="padding:10px 24px;border:2px solid #6c757d;border-radius:8px;background-color:white;color:#343a40;cursor:pointer;font-weight:700;font-size:1.05em;">Sabah</button>
                <input type="hidden" name="template" id="template-input" value="earthlink">
            </div>

            <h2>SOP Metadata</h2>
            <div class="form-grid">
                <div class="form-group">
                    <label>Process Name</label>
                    <input type="text" name="process_name" value="{{ meta.get('process_name', '') }}" class="{{ 'auto-filled' if meta.get('process_name') else '' }}">
                </div>
                <div class="form-group">
                    <label>Process Code</label>
                    <input type="text" name="process_code" value="{{ meta.get('process_code', '') }}" class="{{ 'auto-filled' if meta.get('process_code') else '' }}">
                </div>
                <div class="form-group">
                    <label>Issued By</label>
                    <input type="text" name="issued_by" value="Business Excellence">
                </div>
                <div class="form-group">
                    <label>Release Date</label>
                    <input type="text" name="release_date" id="release_date" placeholder="dd MMM yyyy">
                </div>
                <div class="form-group">
                    <label>Process Owner</label>
                    <input type="text" name="process_owner" value="TBD">
                </div>
            </div>

            <h2>Purpose & Scope</h2>
            <div class="form-grid">
                <div class="form-group full-width">
                    <label>Purpose</label>
                    <textarea name="purpose" rows="2" class="{{ 'auto-filled' if meta.get('purpose') else '' }}">{{ meta.get('purpose', '') }}</textarea>
                </div>
                <div class="form-group full-width">
                    <label>Scope</label>
                    <textarea name="scope" rows="2" class="{{ 'auto-filled' if meta.get('scope') else '' }}">{{ meta.get('scope', '') }}</textarea>
                </div>
            </div>

            <h2>Abbreviations and Definitions</h2>
            <div id="abbreviations-container" class="dynamic-list">
                {% if meta.get('abbreviations_list') %}
                    {% for abbrev in meta['abbreviations_list'] %}
                    <div class="entry">
                        <input type="text" name="abbrev_term[]" placeholder="Term" value="{{ abbrev.get('term', '') }}">
                        <input type="text" name="abbrev_def[]" placeholder="Definition" value="{{ abbrev.get('definition', '') }}">
                        <button type="button" class="remove-btn" onclick="removeEntry(this)">X</button>
                    </div>
                    {% endfor %}
                {% else %}
                    <div class="entry">
                        <input type="text" name="abbrev_term[]" placeholder="Term">
                        <input type="text" name="abbrev_def[]" placeholder="Definition">
                        <button type="button" class="remove-btn" onclick="removeEntry(this)">X</button>
                    </div>
                {% endif %}
            </div>
            <button type="button" class="add-btn" onclick="addEntry('abbreviations-container', 'abbrev_term[]', 'abbrev_def[]', 'Term', 'Definition')">+ Add</button>

            <h2>Referenced Documents and Approvals</h2>
            <div id="references-container" class="dynamic-list">
                {% set refs = [] %}
                {% if meta.get('lane_names') %}
                    {% for lane in meta['lane_names'] %}
                    <div class="entry">
                        <input type="text" name="ref_id[]" placeholder="Document ID" value="N/A">
                        <input type="text" name="ref_title[]" placeholder="Document Title" value="{{ lane }} Approval">
                        <button type="button" class="remove-btn" onclick="removeEntry(this)">X</button>
                    </div>
                    {% endfor %}
                {% endif %}
                {% if meta.get('process_name') or meta.get('process_code') %}
                <div class="entry">
                    <input type="text" name="ref_id[]" placeholder="Document ID" value="DGM- {{ meta.get('process_code', '') }}">
                    <input type="text" name="ref_title[]" placeholder="Document Title" value="{{ meta.get('process_name', '') }} Process Diagram        Notations Meaning">
                    <button type="button" class="remove-btn" onclick="removeEntry(this)">X</button>
                </div>
                {% endif %}
                {% if not meta.get('lane_names') and not meta.get('process_name') %}
                <div class="entry">
                    <input type="text" name="ref_id[]" placeholder="Document ID">
                    <input type="text" name="ref_title[]" placeholder="Document Title">
                    <button type="button" class="remove-btn" onclick="removeEntry(this)">X</button>
                </div>
                {% endif %}
            </div>
            <button type="button" class="add-btn" onclick="addEntry('references-container', 'ref_id[]', 'ref_title[]', 'Document ID', 'Document Title')">+ Add</button>

            <h2>General Policies</h2>
            <div id="policies-container" class="dynamic-list">
                {% if meta.get('general_policies_list') %}
                    {% for policy in meta['general_policies_list'] %}
                    <div class="entry">
                        <input type="text" name="policy_ref[]" placeholder="Ref" value="{{ policy.get('ref', '') }}">
                        <input type="text" name="policy_text[]" placeholder="Policy" value="{{ policy.get('policy', '') }}">
                        <button type="button" class="remove-btn" onclick="removeEntry(this)">X</button>
                    </div>
                    {% endfor %}
                {% else %}
                    <div class="entry">
                        <input type="text" name="policy_ref[]" placeholder="Ref">
                        <input type="text" name="policy_text[]" placeholder="Policy">
                        <button type="button" class="remove-btn" onclick="removeEntry(this)">X</button>
                    </div>
                {% endif %}
            </div>
            <button type="button" class="add-btn" onclick="addEntry('policies-container', 'policy_ref[]', 'policy_text[]', 'Ref', 'Policy')">+ Add</button>

            <div class="btn-row">
                <button type="submit" class="btn btn-primary" id="generate-btn">Generate & Download .docx</button>
                <button type="button" class="btn btn-secondary" onclick="closeModal()">Cancel</button>
            </div>
        </form>
    </div>

    <script>
        // Auto-fill today's date
        (function() {
            var d = new Date();
            var months = ['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec'];
            document.getElementById('release_date').value =
                String(d.getDate()).padStart(2,'0') + ' ' + months[d.getMonth()] + ' ' + d.getFullYear();
        })();

        function selectTemplate(name) {
            document.getElementById('template-input').value = name;
            var all = ['earthlink', 'sana', 'window_world', 'tarabut', 'sabah'];
            all.forEach(function(t) {
                var btn = document.getElementById('btn-' + t);
                if (t === name) {
                    btn.style.backgroundColor = '#007bff';
                    btn.style.color = 'white';
                    btn.style.borderColor = '#007bff';
                } else {
                    btn.style.backgroundColor = 'white';
                    btn.style.color = '#343a40';
                    btn.style.borderColor = '#6c757d';
                }
            });
        }

        function addEntry(containerId, name1, name2, placeholder1, placeholder2) {
            var container = document.getElementById(containerId);
            var div = document.createElement('div');
            div.className = 'entry';
            div.innerHTML =
                '<input type="text" name="' + name1 + '" placeholder="' + placeholder1 + '">' +
                '<input type="text" name="' + name2 + '" placeholder="' + placeholder2 + '">' +
                '<button type="button" class="remove-btn" onclick="removeEntry(this)">X</button>';
            container.appendChild(div);
        }

        function removeEntry(btn) {
            var container = btn.parentElement.parentElement;
            if (container.children.length > 1) {
                btn.parentElement.remove();
            }
        }

        function closeModal() {
            // Tell parent (Camunda Modeler) to close the modal
            if (window.parent && window.parent !== window) {
                window.parent.postMessage('sop-close-modal', '*');
            } else {
                window.close();
            }
        }

        // Handle form submit - send data to parent window via postMessage
        document.getElementById('sop-form').addEventListener('submit', function(e) {
            e.preventDefault();

            var btn = document.getElementById('generate-btn');
            btn.disabled = true;
            btn.innerHTML = '<span class="spinner"></span>Generating...';

            // Serialize form data as URL-encoded string
            var formData = new FormData(this);
            var body = new URLSearchParams(formData).toString();

            // Send to parent (Camunda Modeler) for download
            window.parent.postMessage({
                type: 'sop-generate',
                session_id: '{{ session_id }}',
                body: body
            }, '*');
        });

        // Listen for response from parent
        window.addEventListener('message', function(e) {
            if (e.data === 'sop-download-complete' || e.data === 'sop-download-error') {
                var btn = document.getElementById('generate-btn');
                btn.disabled = false;
                btn.innerHTML = 'Generate & Download .docx';

                if (e.data === 'sop-download-error') {
                    window.alert('Error generating document. Please try again.');
                }
            }
        });
    </script>
</body>
</html>
'''

INDEX_HTML = r'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>BPMN to SOP Generator</title>
    <style>
        body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif; margin: 0; background-color: #f8f9fa; }
        .main-layout { display: flex; gap: 1.5em; padding: 1.5em; max-width: 1400px; margin: 0 auto; }
        .container { flex: 1; background: white; padding: 2em; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.08); }
        .archive-panel { width: 320px; background: white; padding: 1.5em; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.08); max-height: calc(100vh - 3em); overflow-y: auto; }
        h1, h2 { color: #343a40; border-bottom: 1px solid #dee2e6; padding-bottom: 0.5em; margin-top: 1.5em; }
        h1 { margin-top: 0; }
        .form-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 1.5em; }
        .form-group { display: flex; flex-direction: column; }
        label { margin-bottom: 0.5em; font-weight: 600; color: #495057; }
        input[type="text"], textarea { font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px; box-sizing: border-box; transition: border-color 0.2s, box-shadow 0.2s; }
        input[type="text"]:focus, textarea:focus { border-color: #80bdff; outline: 0; box-shadow: 0 0 0 0.2rem rgba(0,123,255,.25); }
        input[type="file"] { border: 1px solid #ced4da; border-radius: 8px; padding: 0.5em; }
        .full-width { grid-column: 1 / -1; }
        .submit-btn { background-color: #007bff; color: white; padding: 0.8em 1.5em; border: none; border-radius: 8px; font-size: 1.1em; font-weight: 600; cursor: pointer; margin-top: 1.5em; transition: background-color 0.2s; }
        .submit-btn:hover { background-color: #0056b3; }
        .submit-btn:disabled { background-color: #5a6268; cursor: not-allowed; }
        #loader { text-align: center; margin-top: 2em; display: none; }
        .spinner { margin: auto; border: 4px solid #f3f3f3; border-top: 4px solid #007bff; border-radius: 50%; width: 40px; height: 40px; animation: spin 1s linear infinite; }
        @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
        .archive-item { background: #f8f9fa; border-radius: 8px; padding: 1em; margin-bottom: 0.75em; }
        .archive-item-header { display: flex; justify-content: space-between; align-items: start; }
        .archive-item-title { font-weight: 600; color: #343a40; font-size: 0.9em; word-break: break-word; }
        .archive-item-date { font-size: 0.75em; color: #6c757d; margin-top: 0.25em; }
        .archive-item-actions { display: flex; gap: 0.5em; margin-top: 0.75em; }
        .archive-btn { padding: 0.4em 0.6em; border: none; border-radius: 4px; cursor: pointer; font-size: 0.75em; }
        .archive-btn-bpmn { background: #17a2b8; color: white; }
        .archive-btn-docx { background: #28a745; color: white; }
        .archive-btn-delete { background: #dc3545; color: white; }
        .archive-btn:hover { opacity: 0.9; }
        .template-btn { padding: 0.7em 2em; border: 2px solid #6c757d; border-radius: 8px; background: white; cursor: pointer; font-weight: 700; font-size: 1.05em; color: #343a40; transition: all 0.2s; }
        .template-btn:hover { border-color: #007bff; color: #007bff; }
        .template-btn.active { background: #007bff; color: white; border-color: #007bff; }
    </style>
</head>
<body>
    <div class="main-layout">
    <div class="container">
        <h1>BPMN to SOP Generator</h1>

        <form id="sop-form" action="/generate" method="post" enctype="multipart/form-data">

        <div id="template-selector" style="display:flex;align-items:center;gap:12px;margin-bottom:20px;padding:14px 16px;background-color:#dee2e6;border-radius:8px;border:2px solid #adb5bd;">
            <span style="margin:0;font-weight:700;color:#212529;font-size:1.1em;">Template:</span>
            <button type="button" id="btn-earthlink" data-template="earthlink" onclick="selectTemplate('earthlink')" style="padding:10px 24px;border:2px solid #007bff;border-radius:8px;background-color:#007bff;color:white;cursor:pointer;font-weight:700;font-size:1.05em;">Earthlink</button>
            <button type="button" id="btn-sana" data-template="sana" onclick="selectTemplate('sana')" style="padding:10px 24px;border:2px solid #6c757d;border-radius:8px;background-color:white;color:#343a40;cursor:pointer;font-weight:700;font-size:1.05em;">SANA</button>
            <button type="button" id="btn-window_world" data-template="window_world" onclick="selectTemplate('window_world')" style="padding:10px 24px;border:2px solid #6c757d;border-radius:8px;background-color:white;color:#343a40;cursor:pointer;font-weight:700;font-size:1.05em;">Window World</button>
            <button type="button" id="btn-tarabut" data-template="tarabut" onclick="selectTemplate('tarabut')" style="padding:10px 24px;border:2px solid #6c757d;border-radius:8px;background-color:white;color:#343a40;cursor:pointer;font-weight:700;font-size:1.05em;">Tarabut</button>
            <button type="button" id="btn-sabah" data-template="sabah" onclick="selectTemplate('sabah')" style="padding:10px 24px;border:2px solid #6c757d;border-radius:8px;background-color:white;color:#343a40;cursor:pointer;font-weight:700;font-size:1.05em;">Sabah</button>
            <input type="hidden" name="template" id="template-input" value="earthlink">
        </div>

        <!-- Collapsible Metadata Section -->
        <button type="button" id="toggle-metadata-btn" onclick="toggleMetadata()" style="background-color: #6c757d; color: white; border: none; border-radius: 8px; padding: 0.6em 1.2em; cursor: pointer; font-weight: 600; font-size: 0.95em; margin-bottom: 1em; width: 100%; text-align: left;">&#9654; Show Metadata &amp; Settings</button>
        <div id="metadata-sections" style="display: none;">

        <!-- History Section -->
        <div style="background-color: #f8f9fa; padding: 1.5em; border-radius: 8px; margin-bottom: 2em;">
            <h3 style="margin-top: 0; color: #495057; font-size: 1.1em;">Load from History</h3>
            <div style="display: grid; grid-template-columns: 1fr auto; gap: 1em; align-items: center;">
                <select id="history-select" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px; background-color: white;">
                    <option value="">-- Select a previous entry --</option>
                </select>
                <button type="button" id="load-history-btn" onclick="loadHistoryEntry()" style="background-color: #17a2b8; color: white; border: none; border-radius: 8px; padding: 0.75em 1.5em; cursor: pointer; font-weight: 600;">Load</button>
            </div>
        </div>

            <h2>SOP Metadata</h2>
            <div class="form-grid">
                <div class="form-group"><label for="process_name">Process Name:</label><input type="text" id="process_name" name="process_name"></div>
                <div class="form-group"><label for="process_code">Process Code:</label><input type="text" id="process_code" name="process_code"></div>
                <div class="form-group"><label for="issued_by">Issued By:</label><input type="text" id="issued_by" name="issued_by" value="Business Excellence"></div>
                <div class="form-group"><label for="release_date">Release Date:</label><input type="text" id="release_date" name="release_date" placeholder="dd MMM yyyy"></div>
                <div class="form-group"><label for="process_owner">Process Owner:</label><input type="text" id="process_owner" name="process_owner" value="TBD"></div>
            </div>

            <h2 style="margin-top: 2em;">SOP Content Sections</h2>
            <div class="form-grid">
                <div class="form-group full-width"><label for="purpose">Purpose:</label><textarea id="purpose" name="purpose" rows="3"></textarea></div>
                <div class="form-group full-width"><label for="scope">Scope:</label><textarea id="scope" name="scope" rows="3"></textarea></div>
            </div>

            <h2 style="margin-top: 2em;">Abbreviations and Definitions</h2>
            <div id="abbreviations-container">
                <div class="abbrev-entry" style="display: grid; grid-template-columns: 1fr 2fr auto; gap: 1em; margin-bottom: 1em;">
                    <input type="text" name="abbrev_term[]" placeholder="Term (e.g., BPMN)" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                    <input type="text" name="abbrev_def[]" placeholder="Definition (e.g., Business Process Model and Notation)" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                    <button type="button" class="remove-btn" onclick="removeAbbrev(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                </div>
            </div>
            <button type="button" onclick="addAbbrev()" style="background-color: #28a745; color: white; border: none; border-radius: 8px; padding: 0.6em 1.2em; cursor: pointer; margin-bottom: 1em;">+ Add Abbreviation</button>

            <h2 style="margin-top: 2em;">Referenced Documents and Approvals</h2>
            <div id="references-container">
                <div class="ref-entry" style="display: grid; grid-template-columns: 1fr 2fr auto; gap: 1em; margin-bottom: 1em;">
                    <input type="text" name="ref_id[]" placeholder="Document ID (e.g., DOC-001)" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                    <input type="text" name="ref_title[]" placeholder="Document Title" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                    <button type="button" class="remove-btn" onclick="removeRef(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                </div>
            </div>
            <button type="button" onclick="addRef()" style="background-color: #28a745; color: white; border: none; border-radius: 8px; padding: 0.6em 1.2em; cursor: pointer; margin-bottom: 1em;">+ Add Referenced Document</button>

            <h2 style="margin-top: 2em;">Ref. General Policies</h2>
            <div id="policies-container">
                <div class="policy-entry" style="display: grid; grid-template-columns: 1fr 2fr auto; gap: 1em; margin-bottom: 1em;">
                    <input type="text" name="policy_ref[]" placeholder="Ref (e.g., 1)" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                    <input type="text" name="policy_text[]" placeholder="General Policy" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                    <button type="button" class="remove-btn" onclick="removePolicy(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                </div>
            </div>
            <button type="button" onclick="addPolicy()" style="background-color: #28a745; color: white; border: none; border-radius: 8px; padding: 0.6em 1.2em; cursor: pointer; margin-bottom: 1em;">+ Add General Policy</button>

            <!-- Process Inputs and Outputs are auto-populated from BPMN start/end events -->

        </div><!-- end metadata-sections -->

            <h2 style="margin-top: 2em;">BPMN or XML Input</h2>
            <div class="form-group">
                <label>Choose Input Type:</label>
                <div>
                    <input type="radio" id="input_type_bpmn" name="input_type" value="bpmn" checked>
                    <label for="input_type_bpmn">Upload BPMN File</label>
                    <input type="radio" id="input_type_xml" name="input_type" value="xml">
                    <label for="input_type_xml">Paste XML Code</label>
                </div>
            </div>

            <div id="bpmn-upload-section">
                <div class="form-group">
                    <label for="bpmn_file">Upload .bpmn (XML) file:</label>
                    <input type="file" id="bpmn_file" name="bpmn_file" accept=".bpmn" required>
                </div>
            </div>

            <div id="xml-paste-section" style="display: none;">
                <div class="form-group full-width">
                    <label for="xml_code">Paste XML Code:</label>
                    <textarea id="xml_code" name="xml_code" rows="10"></textarea>
                </div>
            </div>

            <div class="full-width">
                <button type="submit" class="submit-btn" id="generate-btn">Generate SOP Document</button>
            </div>
        </form>
        <div id="loader">
            <div class="spinner"></div>
            <p><strong>Generating document, please wait...</strong></p>
        </div>

    </div>
            <script>
                function selectTemplate(name) {
                    document.getElementById('template-input').value = name;
                    var all = ['earthlink', 'sana', 'window_world', 'tarabut', 'sabah'];
                    all.forEach(function(t) {
                        var btn = document.getElementById('btn-' + t);
                        if (t === name) {
                            btn.style.backgroundColor = '#007bff';
                            btn.style.color = 'white';
                            btn.style.borderColor = '#007bff';
                        } else {
                            btn.style.backgroundColor = 'white';
                            btn.style.color = '#343a40';
                            btn.style.borderColor = '#6c757d';
                        }
                    });
                }

                function getCookie(name) {
                    let value = "; " + document.cookie;
                    let parts = value.split("; " + name + "=");
                    if (parts.length === 2) return parts.pop().split(";").shift();
                }

                // Auto-fill release date with today's date
                function setTodaysDate() {
                    const today = new Date();
                    const day = String(today.getDate()).padStart(2, '0');
                    const months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec'];
                    const month = months[today.getMonth()];
                    const year = today.getFullYear();
                    const formattedDate = `${day} ${month} ${year}`;
                    document.getElementById('release_date').value = formattedDate;
                }

                // Call setTodaysDate when page loads
                window.addEventListener('DOMContentLoaded', setTodaysDate);

                // History Management Functions
                let historyData = [];

                // Fetch history from server
                async function fetchHistory() {
                    try {
                        const response = await fetch('/api/history');
                        if (response.ok) {
                            historyData = await response.json();
                            populateHistoryDropdown();
                        }
                    } catch (error) {
                        console.error('Error fetching history:', error);
                    }
                }

                // Populate history dropdown
                function populateHistoryDropdown() {
                    const select = document.getElementById('history-select');
                    // Clear existing options except the first one
                    while (select.options.length > 1) {
                        select.remove(1);
                    }

                    // Add history entries
                    historyData.forEach((entry, index) => {
                        const option = document.createElement('option');
                        option.value = index;

                        // Format the display text with date and time
                        const dt = new Date(entry.timestamp);
                        const timestamp = dt.toLocaleDateString() + ' ' + dt.toLocaleTimeString([], {hour: '2-digit', minute:'2-digit'});
                        const processName = entry.process_name || 'Unnamed';
                        const processCode = entry.process_code ? ` (${entry.process_code})` : '';
                        option.textContent = `${timestamp} - ${processName}${processCode}`;

                        select.appendChild(option);
                    });
                }

                // Load selected history entry into form
                function loadHistoryEntry() {
                    const select = document.getElementById('history-select');
                    const selectedIndex = select.value;

                    if (selectedIndex === '') {
                        alert('Please select an entry from history');
                        return;
                    }

                    const entry = historyData[selectedIndex];
                    if (!entry) {
                        alert('Error loading history entry');
                        return;
                    }

                    // Populate form fields
                    document.getElementById('process_name').value = entry.process_name || '';
                    document.getElementById('process_code').value = entry.process_code || '';
                    document.getElementById('purpose').value = entry.purpose || '';
                    document.getElementById('scope').value = entry.scope || '';

                    // Populate abbreviations
                    const abbrevContainer = document.getElementById('abbreviations-container');
                    abbrevContainer.innerHTML = ''; // Clear existing
                    const abbreviations = entry.abbreviations_list || [];
                    if (abbreviations.length > 0) {
                        abbreviations.forEach(abbrev => {
                            const entry = document.createElement('div');
                            entry.style.display = 'grid';
                            entry.style.gridTemplateColumns = '1fr 2fr auto';
                            entry.style.gap = '1em';
                            entry.style.marginBottom = '1em';
                            entry.innerHTML = `
                                <input type="text" name="abbrev_term[]" placeholder="Term" value="${abbrev.term || ''}" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                                <input type="text" name="abbrev_def[]" placeholder="Definition" value="${abbrev.definition || ''}" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                                <button type="button" onclick="removeAbbrev(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                            `;
                            abbrevContainer.appendChild(entry);
                        });
                    } else {
                        addAbbrev(); // Add one empty entry
                    }

                    // Populate referenced documents
                    const refsContainer = document.getElementById('references-container');
                    refsContainer.innerHTML = ''; // Clear existing
                    const references = entry.references_list || [];
                    if (references.length > 0) {
                        references.forEach(ref => {
                            const entry = document.createElement('div');
                            entry.style.display = 'grid';
                            entry.style.gridTemplateColumns = '1fr 2fr auto';
                            entry.style.gap = '1em';
                            entry.style.marginBottom = '1em';
                            entry.innerHTML = `
                                <input type="text" name="ref_id[]" placeholder="Document ID" value="${ref.id || ''}" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                                <input type="text" name="ref_title[]" placeholder="Document Title" value="${ref.title || ''}" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                                <button type="button" onclick="removeRef(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                            `;
                            refsContainer.appendChild(entry);
                        });
                    } else {
                        addRef(); // Add one empty entry
                    }

                    // Populate general policies
                    const policiesContainer = document.getElementById('policies-container');
                    policiesContainer.innerHTML = '';
                    const policies = entry.general_policies_list || [];
                    if (policies.length > 0) {
                        policies.forEach(pol => {
                            const pEntry = document.createElement('div');
                            pEntry.style.display = 'grid';
                            pEntry.style.gridTemplateColumns = '1fr 2fr auto';
                            pEntry.style.gap = '1em';
                            pEntry.style.marginBottom = '1em';
                            pEntry.innerHTML = `
                                <input type="text" name="policy_ref[]" placeholder="Ref" value="${escapeHtml(pol.ref || '')}" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                                <input type="text" name="policy_text[]" placeholder="General Policy" value="${escapeHtml(pol.policy || '')}" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                                <button type="button" onclick="removePolicy(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                            `;
                            policiesContainer.appendChild(pEntry);
                        });
                    } else {
                        addPolicy();
                    }

                    alert('History entry loaded successfully!');
                }

                // Fetch history on page load
                window.addEventListener('DOMContentLoaded', fetchHistory);

                // Dynamic abbreviation entries
                function addAbbrev() {
                    const container = document.getElementById('abbreviations-container');
                    const entry = document.createElement('div');
                    entry.className = 'abbrev-entry';
                    entry.style.display = 'grid';
                    entry.style.gridTemplateColumns = '1fr 2fr auto';
                    entry.style.gap = '1em';
                    entry.style.marginBottom = '1em';
                    entry.innerHTML = `
                        <input type="text" name="abbrev_term[]" placeholder="Term" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                        <input type="text" name="abbrev_def[]" placeholder="Definition" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                        <button type="button" onclick="removeAbbrev(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                    `;
                    container.appendChild(entry);
                }

                function removeAbbrev(btn) {
                    const container = document.getElementById('abbreviations-container');
                    if (container.children.length > 1) {
                        btn.parentElement.remove();
                    } else {
                        alert('At least one abbreviation entry must remain.');
                    }
                }

                // Dynamic reference entries
                function addRef() {
                    const container = document.getElementById('references-container');
                    const entry = document.createElement('div');
                    entry.className = 'ref-entry';
                    entry.style.display = 'grid';
                    entry.style.gridTemplateColumns = '1fr 2fr auto';
                    entry.style.gap = '1em';
                    entry.style.marginBottom = '1em';
                    entry.innerHTML = `
                        <input type="text" name="ref_id[]" placeholder="Document ID" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                        <input type="text" name="ref_title[]" placeholder="Document Title" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                        <button type="button" onclick="removeRef(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                    `;
                    container.appendChild(entry);
                }

                function removeRef(btn) {
                    const container = document.getElementById('references-container');
                    if (container.children.length > 1) {
                        btn.parentElement.remove();
                    } else {
                        alert('At least one reference entry must remain.');
                    }
                }

                // Dynamic general policy entries
                function addPolicy() {
                    const container = document.getElementById('policies-container');
                    const entry = document.createElement('div');
                    entry.className = 'policy-entry';
                    entry.style.display = 'grid';
                    entry.style.gridTemplateColumns = '1fr 2fr auto';
                    entry.style.gap = '1em';
                    entry.style.marginBottom = '1em';
                    entry.innerHTML = `
                        <input type="text" name="policy_ref[]" placeholder="Ref" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                        <input type="text" name="policy_text[]" placeholder="General Policy" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px;">
                        <button type="button" onclick="removePolicy(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                    `;
                    container.appendChild(entry);
                }

                function removePolicy(btn) {
                    const container = document.getElementById('policies-container');
                    if (container.children.length > 1) {
                        btn.parentElement.remove();
                    } else {
                        alert('At least one policy entry must remain.');
                    }
                }

                const inputTypeBpmn = document.getElementById('input_type_bpmn');
                const inputTypeXml = document.getElementById('input_type_xml');
                const bpmnUploadSection = document.getElementById('bpmn-upload-section');
                const xmlPasteSection = document.getElementById('xml-paste-section');
                const bpmnFile = document.getElementById('bpmn_file');
                const xmlCode = document.getElementById('xml_code');

                function toggleInputSections() {
                    if (inputTypeBpmn.checked) {
                        bpmnUploadSection.style.display = 'block';
                        xmlPasteSection.style.display = 'none';
                        bpmnFile.setAttribute('required', 'required');
                        xmlCode.removeAttribute('required');
                    } else {
                        bpmnUploadSection.style.display = 'none';
                        xmlPasteSection.style.display = 'block';
                        xmlCode.setAttribute('required', 'required');
                        bpmnFile.removeAttribute('required');
                    }
                }

                inputTypeBpmn.addEventListener('change', toggleInputSections);
                inputTypeXml.addEventListener('change', toggleInputSections);

                // Initialize on page load
                toggleInputSections();

                document.getElementById('sop-form').addEventListener('submit', async function(e) {
                    e.preventDefault();

                    const btn = document.getElementById('generate-btn');
                    const loader = document.getElementById('loader');
                    const form = e.target;

                    loader.style.display = 'block';
                    btn.disabled = true;
                    btn.textContent = 'Generating...';

                    try {
                        const formData = new FormData(form);
                        const response = await fetch('/generate', {
                            method: 'POST',
                            body: formData
                        });

                        const result = await response.json();

                        if (result.success) {
                            alert('Document saved successfully!\n\n' + result.path);
                            fetchHistory();
                            loadArchives();
                        } else {
                            alert(result.message || 'Generation cancelled or failed');
                        }
                    } catch (error) {
                        console.error('Error:', error);
                        alert('An error occurred during generation');
                    } finally {
                        loader.style.display = 'none';
                        btn.disabled = false;
                        btn.textContent = 'Generate SOP Document';
                    }
                });

                // Archive functions
                async function loadArchives() {
                    try {
                        const response = await fetch('/api/archive/list');
                        const data = await response.json();
                        displayArchives(data.archives || []);
                    } catch (error) {
                        console.error('Error loading archives:', error);
                    }
                }

                let archivesData = [];

                function displayArchives(archives) {
                    archivesData = archives;
                    const select = document.getElementById('archive-select');
                    const actions = document.getElementById('archive-actions');

                    // Clear and populate dropdown
                    select.innerHTML = '<option value="">-- Select a backup --</option>';

                    if (!archives || archives.length === 0) {
                        actions.style.display = 'none';
                        return;
                    }

                    archives.forEach((archive, index) => {
                        const dt = new Date(archive.created_at);
                        const dateTime = dt.toLocaleDateString() + ' ' + dt.toLocaleTimeString([], {hour: '2-digit', minute:'2-digit'});
                        const option = document.createElement('option');
                        option.value = index;
                        option.textContent = `${archive.process_name} - ${dateTime}`;
                        select.appendChild(option);
                    });
                }

                function onArchiveSelect() {
                    const select = document.getElementById('archive-select');
                    const actions = document.getElementById('archive-actions');
                    actions.style.display = select.value !== '' ? 'flex' : 'none';
                }

                function getSelectedArchiveId() {
                    const select = document.getElementById('archive-select');
                    if (select.value === '') return null;
                    return archivesData[parseInt(select.value)].id;
                }

                async function downloadBpmn() {
                    const id = getSelectedArchiveId();
                    if (!id) return;
                    const response = await fetch(`/api/archive/${id}/bpmn`);
                    const result = await response.json();
                    if (result.success) {
                        alert('BPMN saved to: ' + result.path);
                    }
                }

                async function downloadDocx() {
                    const id = getSelectedArchiveId();
                    if (!id) return;
                    const response = await fetch(`/api/archive/${id}/docx`);
                    const result = await response.json();
                    if (result.success) {
                        alert('Document saved to: ' + result.path);
                    }
                }

                async function deleteSelectedArchive() {
                    const id = getSelectedArchiveId();
                    if (!id) return;
                    if (!confirm('Delete this backup?')) return;
                    try {
                        const response = await fetch(`/api/archive/${id}`, { method: 'DELETE' });
                        const result = await response.json();
                        if (result.success) {
                            document.getElementById('archive-actions').style.display = 'none';
                            loadArchives();
                        }
                    } catch (error) {
                        console.error('Error deleting:', error);
                    }
                }

                // Load archives on page load
                window.addEventListener('DOMContentLoaded', loadArchives);

                // --- Toggle Metadata Sections ---
                function toggleMetadata() {
                    const sections = document.getElementById('metadata-sections');
                    const btn = document.getElementById('toggle-metadata-btn');
                    if (sections.style.display === 'none') {
                        sections.style.display = 'block';
                        btn.innerHTML = '&#9660; Hide Metadata &amp; Settings';
                    } else {
                        sections.style.display = 'none';
                        btn.innerHTML = '&#9654; Show Metadata &amp; Settings';
                    }
                }

                // --- BPMN Metadata Auto-Fill ---

                function escapeHtml(text) {
                    const div = document.createElement('div');
                    div.textContent = text;
                    return div.innerHTML;
                }

                function resetFormFields() {
                    // Clear text fields
                    ['process_name', 'process_code', 'purpose', 'scope'].forEach(id => {
                        const field = document.getElementById(id);
                        if (field) {
                            field.value = '';
                            field.style.backgroundColor = '';
                        }
                    });

                    // Reset abbreviations to one empty entry
                    const abbrevContainer = document.getElementById('abbreviations-container');
                    abbrevContainer.innerHTML = '';
                    addAbbrev();

                    // Reset references to one empty entry
                    const refsContainer = document.getElementById('references-container');
                    refsContainer.innerHTML = '';
                    addRef();

                    // Reset policies to one empty entry
                    const policiesContainer = document.getElementById('policies-container');
                    policiesContainer.innerHTML = '';
                    addPolicy();
                }

                function autoFillFromBpmn(metadata) {
                    // Simple text fields - only fill if currently empty
                    const fieldMappings = {
                        'process_name': 'process_name',
                        'process_code': 'process_code',
                        'purpose': 'purpose',
                        'scope': 'scope'
                    };

                    for (const [metaKey, fieldId] of Object.entries(fieldMappings)) {
                        if (metadata[metaKey]) {
                            const field = document.getElementById(fieldId);
                            if (field && !field.value.trim()) {
                                field.value = metadata[metaKey];
                                field.style.backgroundColor = '#e8f5e9';
                                field.addEventListener('input', function() {
                                    this.style.backgroundColor = '';
                                }, { once: true });
                            }
                        }
                    }

                    // Abbreviations - only auto-fill if current entries are all empty
                    if (metadata.abbreviations_list && metadata.abbreviations_list.length > 0) {
                        const container = document.getElementById('abbreviations-container');
                        const existingTerms = container.querySelectorAll('input[name="abbrev_term[]"]');
                        const existingDefs = container.querySelectorAll('input[name="abbrev_def[]"]');

                        let allEmpty = true;
                        existingTerms.forEach((input, i) => {
                            if (input.value.trim() || (existingDefs[i] && existingDefs[i].value.trim())) {
                                allEmpty = false;
                            }
                        });

                        if (allEmpty) {
                            container.innerHTML = '';
                            metadata.abbreviations_list.forEach(abbrev => {
                                const entry = document.createElement('div');
                                entry.className = 'abbrev-entry';
                                entry.style.display = 'grid';
                                entry.style.gridTemplateColumns = '1fr 2fr auto';
                                entry.style.gap = '1em';
                                entry.style.marginBottom = '1em';
                                entry.innerHTML = `
                                    <input type="text" name="abbrev_term[]" placeholder="Term" value="${escapeHtml(abbrev.term || '')}" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px; background-color: #e8f5e9;">
                                    <input type="text" name="abbrev_def[]" placeholder="Definition" value="${escapeHtml(abbrev.definition || '')}" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px; background-color: #e8f5e9;">
                                    <button type="button" onclick="removeAbbrev(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                                `;
                                container.appendChild(entry);
                            });
                        }
                    }

                    // Referenced Documents & Approvals - auto-fill with lane approvals + DMG row
                    if (metadata.lane_names && metadata.lane_names.length > 0) {
                        const refsContainer = document.getElementById('references-container');
                        const existingIds = refsContainer.querySelectorAll('input[name="ref_id[]"]');
                        const existingTitles = refsContainer.querySelectorAll('input[name="ref_title[]"]');

                        let refsAllEmpty = true;
                        existingIds.forEach((input, i) => {
                            if (input.value.trim() || (existingTitles[i] && existingTitles[i].value.trim())) {
                                refsAllEmpty = false;
                            }
                        });

                        if (refsAllEmpty) {
                            refsContainer.innerHTML = '';

                            // Add lane approval rows: N/A | {Lane Name} Approval
                            metadata.lane_names.forEach(laneName => {
                                const entry = document.createElement('div');
                                entry.className = 'ref-entry';
                                entry.style.display = 'grid';
                                entry.style.gridTemplateColumns = '1fr 2fr auto';
                                entry.style.gap = '1em';
                                entry.style.marginBottom = '1em';
                                entry.innerHTML = `
                                    <input type="text" name="ref_id[]" placeholder="Document ID" value="N/A" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px; background-color: #e8f5e9;">
                                    <input type="text" name="ref_title[]" placeholder="Document Title" value="${escapeHtml(laneName)} Approval" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px; background-color: #e8f5e9;">
                                    <button type="button" onclick="removeRef(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                                `;
                                refsContainer.appendChild(entry);
                            });

                            // Add DMG row: DMG-{process_code} | {process_name} Process Diagram        Notations Meaning
                            const pCode = metadata.process_code || '';
                            const pName = metadata.process_name || '';
                            if (pCode || pName) {
                                const dmgId = pCode ? 'DGM- ' + pCode : 'DGM-';
                                const dmgTitle = pName + ' Process Diagram        Notations Meaning';
                                const dmgEntry = document.createElement('div');
                                dmgEntry.className = 'ref-entry';
                                dmgEntry.style.display = 'grid';
                                dmgEntry.style.gridTemplateColumns = '1fr 2fr auto';
                                dmgEntry.style.gap = '1em';
                                dmgEntry.style.marginBottom = '1em';
                                dmgEntry.innerHTML = `
                                    <input type="text" name="ref_id[]" placeholder="Document ID" value="${escapeHtml(dmgId)}" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px; background-color: #e8f5e9;">
                                    <input type="text" name="ref_title[]" placeholder="Document Title" value="${escapeHtml(dmgTitle)}" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px; background-color: #e8f5e9;">
                                    <button type="button" onclick="removeRef(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                                `;
                                refsContainer.appendChild(dmgEntry);
                            }
                        }
                    }

                    // General Policies - auto-fill from BPMN
                    if (metadata.general_policies_list && metadata.general_policies_list.length > 0) {
                        const policiesContainer = document.getElementById('policies-container');
                        const existingRefs = policiesContainer.querySelectorAll('input[name="policy_ref[]"]');
                        const existingTexts = policiesContainer.querySelectorAll('input[name="policy_text[]"]');

                        let policiesAllEmpty = true;
                        existingRefs.forEach((input, i) => {
                            if (input.value.trim() || (existingTexts[i] && existingTexts[i].value.trim())) {
                                policiesAllEmpty = false;
                            }
                        });

                        if (policiesAllEmpty) {
                            policiesContainer.innerHTML = '';
                            metadata.general_policies_list.forEach(policy => {
                                const entry = document.createElement('div');
                                entry.className = 'policy-entry';
                                entry.style.display = 'grid';
                                entry.style.gridTemplateColumns = '1fr 2fr auto';
                                entry.style.gap = '1em';
                                entry.style.marginBottom = '1em';
                                entry.innerHTML = `
                                    <input type="text" name="policy_ref[]" placeholder="Ref" value="${escapeHtml(policy.ref || '')}" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px; background-color: #e8f5e9;">
                                    <input type="text" name="policy_text[]" placeholder="General Policy" value="${escapeHtml(policy.policy || '')}" style="font-size: 1rem; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px; background-color: #e8f5e9;">
                                    <button type="button" onclick="removePolicy(this)" style="background-color: #dc3545; color: white; border: none; border-radius: 8px; padding: 0.75em 1em; cursor: pointer;">Remove</button>
                                `;
                                policiesContainer.appendChild(entry);
                            });
                        }
                    }
                }

                // Auto-fill form from BPMN metadata when file is selected
                document.getElementById('bpmn_file').addEventListener('change', async function(e) {
                    const file = e.target.files[0];
                    if (!file) return;

                    const btn = document.getElementById('generate-btn');
                    const originalText = btn.textContent;
                    btn.textContent = 'Reading BPMN metadata...';
                    btn.disabled = true;

                    try {
                        const formData = new FormData();
                        formData.append('bpmn_file', file);

                        const response = await fetch('/extract-metadata', {
                            method: 'POST',
                            body: formData
                        });

                        if (response.ok) {
                            const data = await response.json();
                            if (data.success && data.metadata) {
                                resetFormFields();
                                autoFillFromBpmn(data.metadata);
                                // Auto-expand the metadata section so user can see filled fields
                                const sections = document.getElementById('metadata-sections');
                                if (sections.style.display === 'none') {
                                    toggleMetadata();
                                }
                            }
                        }
                    } catch (error) {
                        console.error('Error extracting BPMN metadata:', error);
                    } finally {
                        btn.textContent = originalText;
                        btn.disabled = false;
                    }
                });

                // Auto-fill form from BPMN metadata when XML is pasted
                let xmlExtractTimer = null;
                document.getElementById('xml_code').addEventListener('input', function() {
                    // Debounce: wait 800ms after user stops typing/pasting
                    clearTimeout(xmlExtractTimer);
                    const xmlText = this.value.trim();
                    if (!xmlText || xmlText.length < 50) return; // Too short to be valid BPMN
                    xmlExtractTimer = setTimeout(async () => {
                        const btn = document.getElementById('generate-btn');
                        const originalText = btn.textContent;
                        btn.textContent = 'Reading BPMN metadata...';
                        btn.disabled = true;

                        try {
                            const formData = new FormData();
                            formData.append('xml_code', xmlText);

                            const response = await fetch('/extract-metadata', {
                                method: 'POST',
                                body: formData
                            });

                            if (response.ok) {
                                const data = await response.json();
                                if (data.success && data.metadata) {
                                    resetFormFields();
                                    autoFillFromBpmn(data.metadata);
                                    const sections = document.getElementById('metadata-sections');
                                    if (sections.style.display === 'none') {
                                        toggleMetadata();
                                    }
                                }
                            }
                        } catch (error) {
                            console.error('Error extracting BPMN metadata from XML:', error);
                        } finally {
                            btn.textContent = originalText;
                            btn.disabled = false;
                        }
                    }, 800);
                });

            </script>
    </div>

    <!-- Archive Panel on Right -->
    <div class="archive-panel">
        <h3 style="margin-top: 0; color: #343a40; border-bottom: 1px solid #dee2e6; padding-bottom: 0.5em;">Backups</h3>
        <select id="archive-select" onchange="onArchiveSelect()" style="width: 100%; padding: 0.75em; border: 1px solid #ced4da; border-radius: 8px; font-size: 0.9em; background: white;">
            <option value="">-- Select a backup --</option>
        </select>
        <div id="archive-actions" style="display: none; margin-top: 1em; gap: 0.5em;">
            <button class="archive-btn archive-btn-bpmn" onclick="downloadBpmn()" style="flex: 1; padding: 0.6em;">BPMN</button>
            <button class="archive-btn archive-btn-docx" onclick="downloadDocx()" style="flex: 1; padding: 0.6em;">Word</button>
            <button class="archive-btn archive-btn-delete" onclick="deleteSelectedArchive()" style="padding: 0.6em 1em;">&#128465;</button>
        </div>
    </div>

    </div>
</body>
</html>
'''

# --- App Data Directory (store runtime data in AppData, not exe folder) ---
def get_app_data_dir():
    if os.name == 'nt':  # Windows
        base = os.environ.get('LOCALAPPDATA', os.path.expanduser('~'))
    else:  # Linux/Mac
        base = os.path.join(os.path.expanduser('~'), '.config')
    app_dir = os.path.join(base, 'SOP_Generator')
    os.makedirs(app_dir, exist_ok=True)
    return app_dir

APP_DATA_DIR = get_app_data_dir()

# --- History Manager ---
history_manager = HistoryManager(os.path.join(APP_DATA_DIR, 'history'))

# --- Archive Manager ---
archive_manager = ArchiveManager(os.path.join(APP_DATA_DIR, 'archives'), os.path.join(APP_DATA_DIR, 'archive.db'))

# --- Core Logic using our custom BPMN parser ---
def parse_bpmn_to_context(xml_content, metadata):
    """
    Wrapper function that calls our comprehensive BPMN parser
    """
    return parse_bpmn_to_sop(xml_content, metadata)


def create_word_doc_from_template(context, template_name='earthlink'):
    """
    Create Word document with multi-paragraph structure and precise formatting
    Following Guideline V2 specifications
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        template_map = {
            'sana': 'sana_template.docx',
            'window_world': 'window_world_template.docx',
            'tarabut': 'tarabut_template.docx',
            'sabah': 'sabah_template.docx',
        }
        template_file = template_map.get(template_name, 'final_master_template_2.docx')
        template_path = resource_path(template_file)

        # First, render metadata using docxtpl (for any {{variables}} in headers/etc)
        doc_template = DocxTemplate(template_path)

        # Create a simple context for metadata rendering (excluding steps)
        metadata_context = {k: v for k, v in context.items() if k != 'steps'}
        doc_template.render(metadata_context)

        # Now work with the rendered document using python-docx
        # Save to memory and reload as Document for easier manipulation
        temp_stream = io.BytesIO()
        doc_template.save(temp_stream)
        temp_stream.seek(0)
        doc = Document(temp_stream)

        # Get the tables
        if not doc.tables:
            raise Exception("No tables found in template")

        # Fix font for all front matter tables to size 12
        # Table 0: Header (Process Name, Code, etc.)
        # Table 1: Purpose
        # Table 2: Scope
        # Table 3: Abbreviations and Definitions
        # Table 4: Referenced Documents and Approvals
        # Table 5: Key Process Inputs/Outputs
        for table_idx in [0, 1, 2, 3, 4, 5]:
            if len(doc.tables) > table_idx:
                table = doc.tables[table_idx]
                # Iterate through all cells in the table
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Avenir LT Std 45 Book'
                                run.font.size = Pt(12)

        # Populate Table 4: Abbreviations and Definitions (index 3)
        if len(doc.tables) > 3:
            abbrev_table = doc.tables[3]
            abbreviations = context.get('abbreviations_list', [])

            # Clear existing rows (keep row 0=title and row 1=headers)
            rows_to_delete = list(range(len(abbrev_table.rows) - 1, 1, -1))
            for row_idx in rows_to_delete:
                abbrev_table._element.remove(abbrev_table.rows[row_idx]._element)

            # Add rows for each abbreviation
            if abbreviations:
                for abbrev in abbreviations:
                    row = abbrev_table.add_row()
                    row.cells[0].text = abbrev.get('term', '')
                    row.cells[1].text = abbrev.get('definition', '')
                    # Format cells - Terms column (column 0) should be bold
                    for idx, cell in enumerate(row.cells):
                        para = cell.paragraphs[0]
                        if para.runs:
                            para.runs[0].font.name = 'Avenir LT Std 45 Book'
                            para.runs[0].font.size = Pt(12)  # Changed from 11 to 12
                            # Make Terms column (column 0) bold
                            if idx == 0:
                                para.runs[0].font.bold = True
            else:
                # Add one empty row if no abbreviations provided
                row = abbrev_table.add_row()
                row.cells[0].text = 'N/A'
                row.cells[1].text = 'N/A'
                # Format cells
                for cell in row.cells:
                    para = cell.paragraphs[0]
                    if para.runs:
                        para.runs[0].font.name = 'Avenir LT Std 45 Book'
                        para.runs[0].font.size = Pt(12)

        # Populate Table 5: Referenced Documents and Approvals (index 4)
        if len(doc.tables) > 4:
            ref_table = doc.tables[4]
            references = context.get('references_list', [])

            # Clear existing rows (keep row 0=title and row 1=headers)
            rows_to_delete = list(range(len(ref_table.rows) - 1, 1, -1))
            for row_idx in rows_to_delete:
                ref_table._element.remove(ref_table.rows[row_idx]._element)

            # Add rows for each reference
            if references:
                for ref in references:
                    row = ref_table.add_row()
                    row.cells[0].text = ref.get('id', '')
                    row.cells[1].text = ref.get('title', '')
                    # Format cells - Document ID column (column 0) should be bold
                    for idx, cell in enumerate(row.cells):
                        para = cell.paragraphs[0]
                        if para.runs:
                            para.runs[0].font.name = 'Avenir LT Std 45 Book'
                            para.runs[0].font.size = Pt(12)  # Changed from 11 to 12
                            # Make Document ID column (column 0) bold
                            if idx == 0:
                                para.runs[0].font.bold = True
            else:
                # Add one empty row if no references provided
                row = ref_table.add_row()
                row.cells[0].text = 'N/A'
                row.cells[1].text = 'N/A'
                # Format cells
                for cell in row.cells:
                    para = cell.paragraphs[0]
                    if para.runs:
                        para.runs[0].font.name = 'Avenir LT Std 45 Book'
                        para.runs[0].font.size = Pt(12)

        # Get the process description table (Table 6 - Table 7 is General Policies)
        table = doc.tables[6]

        # Clear all rows except header (row 0)
        rows_to_delete = list(range(len(table.rows) - 1, 0, -1))  # Delete in reverse order
        for row_idx in rows_to_delete:
            table._element.remove(table.rows[row_idx]._element)

        # Now add rows for each step with proper formatting
        steps = context.get('steps', [])

        for step in steps:
            # Add new row
            new_row = table.add_row()
            cells = new_row.cells

            # --- Cell 0: Ref number ---
            ref_cell = cells[0]
            ref_cell.text = ''  # Clear any existing text
            ref_para = ref_cell.paragraphs[0]
            ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if step['ref']:  # Only add ref if it exists
                run = ref_para.add_run(step['ref'])
                run.font.name = 'Avenir LT Std 45 Book'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)  # RED

            # --- Cell 1: Process Description (multi-paragraph) ---
            desc_cell = cells[1]
            desc_cell.text = ''  # Clear any existing text

            # Remove default paragraph
            if desc_cell.paragraphs:
                p = desc_cell.paragraphs[0]
                p._element.getparent().remove(p._element)

            # Add each paragraph with specific formatting
            for para_data in step['paragraphs']:
                para = desc_cell.add_paragraph()

                # Set alignment
                if para_data['alignment'] == 'CENTER':
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif para_data['alignment'] == 'JUSTIFY':
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Add text with formatting
                if para_data['text']:  # Only add run if there's text
                    run = para.add_run(para_data['text'])
                    run.font.name = 'Avenir LT Std 45 Book'
                    run.font.size = Pt(para_data['font_size'])
                    run.font.bold = para_data['bold']
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black

            # --- Cells 2-6: RACI + SLA ---
            # All RACI fields must be Avenir LT Std 45 Book font size 9
            raci = step.get('raci', {})
            raci_map = {2: 'responsible', 3: 'accountable', 4: 'consulted', 5: 'informed'}
            for i in range(2, 7):
                # Clear cell and get paragraph
                cells[i].text = ''
                para = cells[i].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Set paragraph style font (this ensures the cell has the right font even when empty)
                para.style.font.name = 'Avenir LT Std 45 Book'
                para.style.font.size = Pt(9)

                if i == 6:  # SLA column - leave blank (handled by SLA merge logic)
                    run = para.add_run('')
                    run.font.name = 'Avenir LT Std 45 Book'
                    run.font.size = Pt(9)
                elif i in raci_map:  # R, A, C, I columns - use lane RACI values
                    value = raci.get(raci_map[i], 'N/A') or 'N/A'
                    run = para.add_run(value)
                    run.font.name = 'Avenir LT Std 45 Book'
                    run.font.size = Pt(9)

            # --- Apply gateway shading if needed ---
            if step.get('is_gateway', False):
                for cell_idx, cell in enumerate(new_row.cells):
                    if cell_idx == 6:
                        continue  # Skip SLA column - handled separately
                    tcPr = cell._element.get_or_add_tcPr()
                    # Remove existing shading
                    for shd in tcPr.findall(qn('w:shd')):
                        tcPr.remove(shd)
                    # Apply D9D9D9 shading
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'D9D9D9')  # 15% darker gray
                    shd.set(qn('w:val'), 'clear')
                    tcPr.append(shd)

        # --- SLA shading and vertical merging ---
        # Compute merge ranges: (start_step_idx, end_step_idx, sla_value)
        sla_merges = []
        idx = 0
        while idx < len(steps):
            step_item = steps[idx]
            sla = step_item.get('sla')
            sla_group = step_item.get('sla_group')

            if sla and not sla_group:
                # Task with own SLA - include following gateway cases
                merge_start = idx
                merge_end = idx
                j = idx + 1
                while j < len(steps) and steps[j].get('is_gateway', False):
                    merge_end = j
                    j += 1
                sla_merges.append((merge_start, merge_end, sla))
                idx = j
            elif sla_group:
                # Task in SLA group - include all group members + their gateway cases
                merge_start = idx
                merge_end = idx
                group_sla = sla
                j = idx + 1
                # Include gateway cases of this task
                while j < len(steps) and steps[j].get('is_gateway', False):
                    merge_end = j
                    j += 1
                # Continue with next tasks in same group
                while j < len(steps) and steps[j].get('sla_group') == sla_group:
                    merge_end = j
                    j += 1
                    # Include their gateway cases too
                    while j < len(steps) and steps[j].get('is_gateway', False):
                        merge_end = j
                        j += 1
                sla_merges.append((merge_start, merge_end, group_sla))
                idx = j
            else:
                # No SLA - skip this step and any following gateway cases
                idx += 1
                while idx < len(steps) and steps[idx].get('is_gateway', False):
                    idx += 1

        # Apply SLA shading and vertical merge
        for merge_start, merge_end, sla_value in sla_merges:
            for row_offset in range(merge_start, merge_end + 1):
                table_row_idx = row_offset + 1  # Row 0 is header
                if table_row_idx >= len(table.rows):
                    break
                sla_cell = table.rows[table_row_idx].cells[6]
                tcPr = sla_cell._element.get_or_add_tcPr()

                # Apply F2F2F2 shading (White, Background 1, Darker 5%)
                for existing_shd in tcPr.findall(qn('w:shd')):
                    tcPr.remove(existing_shd)
                shd_elem = OxmlElement('w:shd')
                shd_elem.set(qn('w:fill'), 'F2F2F2')
                shd_elem.set(qn('w:val'), 'clear')
                tcPr.append(shd_elem)

                # Vertical merge if multiple rows in range
                if merge_end > merge_start:
                    vMerge = OxmlElement('w:vMerge')
                    if row_offset == merge_start:
                        vMerge.set(qn('w:val'), 'restart')
                    tcPr.append(vMerge)

                # Write SLA value only in first row of merge
                if row_offset == merge_start:
                    p = sla_cell.paragraphs[0]
                    if p.runs:
                        p.runs[0].text = sla_value

        # --- Populate Table 7: General Policies ---
        if len(doc.tables) > 7:
            policies_table = doc.tables[7]
            policies = context.get('general_policies_list', [])

            # Clear existing data rows (keep row 0=headers only)
            rows_to_delete = list(range(len(policies_table.rows) - 1, 0, -1))
            for row_idx in rows_to_delete:
                policies_table._element.remove(policies_table.rows[row_idx]._element)

            if policies:
                for idx, policy in enumerate(policies, start=1):
                    row = policies_table.add_row()
                    # Ref column: same style as Process Description Ref (Avenir, 14, bold, red, centered)
                    ref_cell = row.cells[0]
                    ref_cell.text = ''
                    ref_para = ref_cell.paragraphs[0]
                    ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = ref_para.add_run(str(idx))
                    run.font.name = 'Avenir LT Std 45 Book'
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    # Policy text column: Avenir, 12, bold
                    text_cell = row.cells[1]
                    text_cell.text = ''
                    text_para = text_cell.paragraphs[0]
                    run = text_para.add_run(policy.get('policy', ''))
                    run.font.name = 'Avenir LT Std 45 Book'
                    run.font.size = Pt(12)
                    run.font.bold = True
            else:
                row = policies_table.add_row()
                # N/A ref cell
                ref_cell = row.cells[0]
                ref_cell.text = ''
                ref_para = ref_cell.paragraphs[0]
                ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = ref_para.add_run('N/A')
                run.font.name = 'Avenir LT Std 45 Book'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                # N/A policy text cell
                text_cell = row.cells[1]
                text_cell.text = ''
                text_para = text_cell.paragraphs[0]
                run = text_para.add_run('N/A')
                run.font.name = 'Avenir LT Std 45 Book'
                run.font.size = Pt(12)
                run.font.bold = True

        # Save to BytesIO
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    except Exception as e:
        print(f"[ERROR] Word Doc Generation Failed: {e}")
        import traceback
        traceback.print_exc()
        return None

@app.route('/')
def index():
    try:
        debug_log(f"index() called, template_folder={app.template_folder}")
        return render_template_string(INDEX_HTML)
    except Exception as e:
        import traceback
        debug_log(f"index() ERROR: {e}")
        debug_log(traceback.format_exc())
        raise

@app.route('/extract-metadata', methods=['POST'])
def extract_metadata():
    """Extract metadata from uploaded BPMN file or pasted XML for form auto-population"""
    bpmn_content = None

    # Check for file upload first
    if 'bpmn_file' in request.files:
        file = request.files['bpmn_file']
        if file.filename != '':
            bpmn_content = file.read()

    # Fall back to raw XML text
    if not bpmn_content:
        xml_code = request.form.get('xml_code', '').strip()
        if xml_code:
            bpmn_content = xml_code.encode('utf-8')

    if not bpmn_content:
        return jsonify({'error': 'No BPMN content provided'}), 400

    try:
        metadata = extract_metadata_from_bpmn(bpmn_content)
        return jsonify({'success': True, 'metadata': metadata})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/history', methods=['GET'])
def get_history():
    """Get all history entries"""
    history_manager.set_user('local')
    history = history_manager.get_all()
    return jsonify(history)

@app.route('/api/history/<int:index>', methods=['GET'])
def get_history_entry(index):
    """Get a specific history entry by index"""
    history_manager.set_user('local')
    entry = history_manager.get_entry(index)
    if entry:
        return jsonify(entry)
    return jsonify({'error': 'Entry not found'}), 404

# --- Archive API Endpoints ---

@app.route('/api/user/set', methods=['POST'])
def set_user():
    """Set current user ID in session"""
    data = request.get_json()
    user_id = data.get('user_id', '').strip().lower()  # Normalize to lowercase

    if not user_id:
        return jsonify({'error': 'User ID required'}), 400

    # Simple validation - alphanumeric and underscores only
    if not re.match(r'^[a-zA-Z0-9_]+$', user_id):
        return jsonify({'error': 'User ID can only contain letters, numbers, and underscores'}), 400

    session['user_id'] = user_id
    return jsonify({'success': True, 'user_id': user_id})

@app.route('/api/user/get', methods=['GET'])
def get_user():
    """Get current user ID from session"""
    user_id = session.get('user_id', None)
    return jsonify({'user_id': user_id})

@app.route('/api/archive/save', methods=['POST'])
def save_archive():
    """Save current BPMN and Word files to archive"""
    user_id = session.get('user_id', None)
    if not user_id:
        return jsonify({'error': 'No user logged in'}), 401

    # Get files from request
    if 'bpmn_file' not in request.files or 'docx_file' not in request.files:
        return jsonify({'error': 'Both BPMN and Word files required'}), 400

    bpmn_file = request.files['bpmn_file']
    docx_file = request.files['docx_file']
    process_name = request.form.get('process_name', 'Untitled Process')

    if bpmn_file.filename == '' or docx_file.filename == '':
        return jsonify({'error': 'Empty file names'}), 400

    # Save files temporarily
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix='.bpmn') as tmp_bpmn:
        bpmn_file.save(tmp_bpmn.name)
        bpmn_temp_path = tmp_bpmn.name

    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
        docx_file.save(tmp_docx.name)
        docx_temp_path = tmp_docx.name

    try:
        # Save to archive
        archive_id = archive_manager.save_archive(
            user_id=user_id,
            process_name=process_name,
            bpmn_file_path=bpmn_temp_path,
            docx_file_path=docx_temp_path
        )

        return jsonify({
            'success': True,
            'archive_id': archive_id,
            'message': 'Files archived successfully'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # Clean up temp files
        try:
            os.remove(bpmn_temp_path)
            os.remove(docx_temp_path)
        except:
            pass

@app.route('/api/archive/list', methods=['GET'])
def list_archives():
    """Get all archives"""
    archives = archive_manager.get_user_archives('local')
    return jsonify({'archives': archives})

@app.route('/api/archive/<int:archive_id>/bpmn', methods=['GET'])
def download_archive_bpmn(archive_id):
    """Download archived BPMN file"""
    archive = archive_manager.get_archive(archive_id)
    if not archive:
        return jsonify({'error': 'Archive not found'}), 404

    bpmn_path = archive_manager.get_file_path(archive_id, 'bpmn')
    if not bpmn_path:
        return jsonify({'error': 'BPMN file not found'}), 404

    # Use save dialog for download
    import tkinter as tk
    from tkinter import filedialog
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    save_path = filedialog.asksaveasfilename(
        parent=root,
        defaultextension='.bpmn',
        filetypes=[('BPMN File', '*.bpmn')],
        initialfile=archive['bpmn_filename'],
        title='Save BPMN File'
    )
    root.destroy()

    if save_path:
        import shutil
        shutil.copy(bpmn_path, save_path)
        return jsonify({'success': True, 'path': save_path})
    return jsonify({'success': False, 'message': 'Cancelled'})

@app.route('/api/archive/<int:archive_id>/docx', methods=['GET'])
def download_archive_docx(archive_id):
    """Download archived Word file"""
    archive = archive_manager.get_archive(archive_id)
    if not archive:
        return jsonify({'error': 'Archive not found'}), 404

    docx_path = archive_manager.get_file_path(archive_id, 'docx')
    if not docx_path:
        return jsonify({'error': 'Word file not found'}), 404

    # Use save dialog for download
    import tkinter as tk
    from tkinter import filedialog
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    save_path = filedialog.asksaveasfilename(
        parent=root,
        defaultextension='.docx',
        filetypes=[('Word Document', '*.docx')],
        initialfile=archive['docx_filename'],
        title='Save Word Document'
    )
    root.destroy()

    if save_path:
        import shutil
        shutil.copy(docx_path, save_path)
        return jsonify({'success': True, 'path': save_path})
    return jsonify({'success': False, 'message': 'Cancelled'})

@app.route('/api/archive/<int:archive_id>', methods=['DELETE'])
def delete_archive(archive_id):
    """Delete an archive"""
    success = archive_manager.delete_archive(archive_id, 'local')
    if success:
        return jsonify({'success': True, 'message': 'Archive deleted'})
    return jsonify({'error': 'Archive not found'}), 404

@app.route('/api/generate-from-xml', methods=['POST'])
def api_generate_from_xml():
    """API endpoint for Camunda Modeler plugin - accepts XML, returns .docx"""
    try:
        data = request.get_json()
        if not data or 'xml' not in data:
            return jsonify({'error': 'No XML provided'}), 400

        bpmn_content = data['xml'].encode('utf-8')
        metadata = data.get('metadata', {})

        # Extract BPMN metadata for fields not provided
        bpmn_metadata = extract_metadata_from_bpmn(bpmn_content)
        for field in ['process_name', 'process_code', 'purpose', 'scope']:
            if not metadata.get(field, '').strip() and field in bpmn_metadata:
                metadata[field] = bpmn_metadata[field]

        # Auto-populate abbreviations if not provided
        if 'abbreviations_list' not in metadata:
            metadata['abbreviations_list'] = bpmn_metadata.get('abbreviations_list', [])

        # Auto-populate references if not provided
        if 'references_list' not in metadata:
            references = []
            lane_names = bpmn_metadata.get('lane_names', [])
            for lane_name in lane_names:
                references.append({'id': 'N/A', 'title': f"{lane_name} Approval"})
            process_code = metadata.get('process_code', bpmn_metadata.get('process_code', ''))
            process_name = metadata.get('process_name', bpmn_metadata.get('process_name', ''))
            if process_code or process_name:
                diagram_id = f"DGM- {process_code}" if process_code else "DGM-"
                diagram_title = f"{process_name} Process Diagram        Notations Meaning"
                references.append({'id': diagram_id, 'title': diagram_title})
            metadata['references_list'] = references

        # Auto-populate general policies if not provided
        if 'general_policies_list' not in metadata:
            metadata['general_policies_list'] = bpmn_metadata.get('general_policies_list', [])

        # Get template selection
        template_name = data.get('template', 'earthlink')

        context = parse_bpmn_to_context(bpmn_content, metadata)
        file_stream = create_word_doc_from_template(context, template_name=template_name)

        if not file_stream:
            return jsonify({'error': 'Failed to generate document'}), 500

        file_stream.seek(0)
        output_name = metadata.get('process_name', bpmn_metadata.get('process_name', 'SOP_Document'))
        response = make_response(file_stream.read())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{output_name}.docx"'
        response.headers['Access-Control-Allow-Origin'] = '*'
        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

import uuid as _uuid
_preview_sessions = {}

@app.route('/api/upload-xml', methods=['POST', 'OPTIONS'])
def api_upload_xml():
    """Store XML temporarily and return a session ID for the preview page"""
    if request.method == 'OPTIONS':
        response = make_response()
        response.headers['Access-Control-Allow-Origin'] = '*'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return response

    data = request.get_json()
    if not data or 'xml' not in data:
        return jsonify({'error': 'No XML provided'}), 400

    session_id = str(_uuid.uuid4())
    bpmn_content = data['xml'].encode('utf-8')
    metadata = extract_metadata_from_bpmn(bpmn_content)

    _preview_sessions[session_id] = {
        'xml': data['xml'],
        'metadata': metadata
    }

    resp = jsonify({'session_id': session_id, 'metadata': metadata})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

@app.route('/preview/<session_id>')
def preview_page(session_id):
    """Serve the SOP preview/edit form pre-populated from BPMN"""
    session_data = _preview_sessions.get(session_id)
    if not session_data:
        return 'Session expired. Please try again from Camunda Modeler.', 404

    meta = session_data['metadata']
    return render_template_string(PREVIEW_HTML, session_id=session_id, meta=meta)

@app.route('/api/generate-and-download/<session_id>', methods=['POST', 'OPTIONS'])
def api_generate_and_download(session_id):
    """Generate .docx from stored XML + user-edited metadata, return binary"""
    if request.method == 'OPTIONS':
        response = make_response()
        response.headers['Access-Control-Allow-Origin'] = '*'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return response

    session_data = _preview_sessions.get(session_id)
    if not session_data:
        resp = jsonify({'error': 'Session expired'})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 404

    try:
        bpmn_content = session_data['xml'].encode('utf-8')
        form_data = request.form.to_dict()

        metadata = {}
        for field in ['process_name', 'process_code', 'issued_by', 'release_date', 'process_owner', 'purpose', 'scope']:
            metadata[field] = form_data.get(field, '').strip()

        # Fall back to BPMN metadata for empty fields
        bpmn_meta = session_data['metadata']
        for field in ['process_name', 'process_code', 'purpose', 'scope']:
            if not metadata.get(field) and field in bpmn_meta:
                metadata[field] = bpmn_meta[field]

        # Parse abbreviations
        abbrev_terms = request.form.getlist('abbrev_term[]')
        abbrev_defs = request.form.getlist('abbrev_def[]')
        abbreviations = []
        for term, definition in zip(abbrev_terms, abbrev_defs):
            if term.strip() or definition.strip():
                abbreviations.append({'term': term.strip(), 'definition': definition.strip()})
        if not abbreviations and 'abbreviations_list' in bpmn_meta:
            abbreviations = bpmn_meta['abbreviations_list']
        metadata['abbreviations_list'] = abbreviations

        # Parse references
        ref_ids = request.form.getlist('ref_id[]')
        ref_titles = request.form.getlist('ref_title[]')
        references = []
        for doc_id, title in zip(ref_ids, ref_titles):
            if doc_id.strip() or title.strip():
                references.append({'id': doc_id.strip(), 'title': title.strip()})
        if not references:
            lane_names = bpmn_meta.get('lane_names', [])
            for lane_name in lane_names:
                references.append({'id': 'N/A', 'title': f"{lane_name} Approval"})
            process_code = metadata.get('process_code', '')
            process_name = metadata.get('process_name', '')
            if process_code or process_name:
                references.append({
                    'id': f"DGM- {process_code}" if process_code else "DGM-",
                    'title': f"{process_name} Process Diagram        Notations Meaning"
                })
        metadata['references_list'] = references

        # Parse policies
        policy_refs = request.form.getlist('policy_ref[]')
        policy_texts = request.form.getlist('policy_text[]')
        policies = []
        for ref, text in zip(policy_refs, policy_texts):
            if ref.strip() or text.strip():
                policies.append({'ref': ref.strip(), 'policy': text.strip()})
        metadata['general_policies_list'] = policies

        # Get template selection
        template_name = form_data.get('template', 'earthlink')

        context = parse_bpmn_to_context(bpmn_content, metadata)
        file_stream = create_word_doc_from_template(context, template_name=template_name)

        if not file_stream:
            return jsonify({'error': 'Failed to generate document'}), 500

        file_stream.seek(0)
        output_name = metadata.get('process_name', 'SOP_Document')

        # Clean up session
        _preview_sessions.pop(session_id, None)

        response = make_response(file_stream.read())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{output_name}.docx"'
        response.headers['Access-Control-Allow-Origin'] = '*'
        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        resp = jsonify({'error': str(e)})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 500

@app.route('/api/generate-from-xml', methods=['OPTIONS'])
def api_generate_from_xml_options():
    """CORS preflight for Camunda Modeler plugin"""
    response = make_response()
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    return response

@app.route('/generate', methods=['POST'])
def generate_sop():
    input_type = request.form.get('input_type')
    bpmn_content = None
    output_name = None

    if input_type == 'bpmn':
        if 'bpmn_file' not in request.files:
            return "No file part", 400
        file = request.files['bpmn_file']
        if file.filename == '':
            return "No selected file", 400
        if file:
            bpmn_content = file.read()
            # Use BPMN filename (without extension) for output
            output_name = file.filename.rsplit('.', 1)[0] if '.' in file.filename else file.filename
    elif input_type == 'xml':
        xml_code = request.form.get('xml_code')
        if not xml_code:
            return "No XML code provided", 400
        bpmn_content = xml_code.encode('utf-8') # Encode to bytes for consistency with file.read()

        # Extract pool/process name from XML
        try:
            from lxml import etree
            root = etree.fromstring(bpmn_content)
            ns = {'bpmn': 'http://www.omg.org/spec/BPMN/20100524/MODEL'}

            # Try to get participant (pool) name
            participant = root.find('.//bpmn:participant', namespaces=ns)
            if participant is not None and participant.get('name'):
                output_name = participant.get('name')
            else:
                # Fallback to process name
                process = root.find('.//bpmn:process', namespaces=ns)
                if process is not None and process.get('name'):
                    output_name = process.get('name')
        except:
            pass
    else:
        return "Invalid input type selected", 400

    if bpmn_content:
        metadata = request.form.to_dict()

        # Extract BPMN metadata for fields user left empty
        bpmn_metadata = extract_metadata_from_bpmn(bpmn_content)

        # For simple string fields: use form value if non-empty, else BPMN value
        for field in ['process_name', 'process_code', 'purpose', 'scope']:
            form_value = metadata.get(field, '').strip()
            if not form_value and field in bpmn_metadata:
                metadata[field] = bpmn_metadata[field]

        # Parse abbreviation entries
        abbrev_terms = request.form.getlist('abbrev_term[]')
        abbrev_defs = request.form.getlist('abbrev_def[]')
        abbreviations = []
        for term, definition in zip(abbrev_terms, abbrev_defs):
            if term.strip() or definition.strip():  # Only add non-empty entries
                abbreviations.append({'term': term.strip(), 'definition': definition.strip()})
        # If no user abbreviations, fall back to BPMN-extracted ones
        if not abbreviations and 'abbreviations_list' in bpmn_metadata:
            abbreviations = bpmn_metadata['abbreviations_list']
        metadata['abbreviations_list'] = abbreviations

        # Parse reference document entries
        ref_ids = request.form.getlist('ref_id[]')
        ref_titles = request.form.getlist('ref_title[]')
        references = []
        for doc_id, title in zip(ref_ids, ref_titles):
            if doc_id.strip() or title.strip():  # Only add non-empty entries
                references.append({'id': doc_id.strip(), 'title': title.strip()})

        # Auto-add lane approvals + DGM row ONLY if form sent no references
        # (i.e., JavaScript auto-fill didn't run). If user already has references
        # from auto-fill (possibly edited), don't add duplicates.
        if not references:
            lane_names = bpmn_metadata.get('lane_names', [])
            for lane_name in lane_names:
                references.append({'id': 'N/A', 'title': f"{lane_name} Approval"})

            process_code = metadata.get('process_code', '').strip()
            if not process_code:
                process_code = bpmn_metadata.get('process_code', '')
            process_name = metadata.get('process_name', '').strip()
            if not process_name:
                process_name = bpmn_metadata.get('process_name', '')
            if process_code or process_name:
                diagram_id = f"DGM- {process_code}" if process_code else "DGM-"
                diagram_title = f"{process_name} Process Diagram        Notations Meaning"
                references.append({'id': diagram_id, 'title': diagram_title})

        metadata['references_list'] = references

        # Parse general policy entries
        policy_refs = request.form.getlist('policy_ref[]')
        policy_texts = request.form.getlist('policy_text[]')
        policies = []
        for ref, text in zip(policy_refs, policy_texts):
            if ref.strip() or text.strip():
                policies.append({'ref': ref.strip(), 'policy': text.strip()})
        metadata['general_policies_list'] = policies

        # Get template selection
        template_name = metadata.get('template', 'earthlink')

        context = parse_bpmn_to_context(bpmn_content, metadata)
        file_stream = create_word_doc_from_template(context, template_name=template_name)
        if file_stream:
            # Save to history on successful generation
            history_manager.set_user('local')
            history_data = {
                'process_name': metadata.get('process_name', ''),
                'process_code': metadata.get('process_code', ''),
                'purpose': metadata.get('purpose', ''),
                'scope': metadata.get('scope', ''),
                'abbreviations_list': metadata.get('abbreviations_list', []),
                'references_list': metadata.get('references_list', []),
                'general_policies_list': metadata.get('general_policies_list', [])
            }
            history_manager.add_entry(history_data)

            # Use extracted name or fallback
            if not output_name:
                output_name = metadata.get('process_name', 'Generated')

            # Show save dialog and save file
            import tkinter as tk
            from tkinter import filedialog

            # Create hidden root window for dialog
            root = tk.Tk()
            root.withdraw()
            root.attributes('-topmost', True)

            # Show save dialog
            save_path = filedialog.asksaveasfilename(
                parent=root,
                defaultextension='.docx',
                filetypes=[('Word Document', '*.docx')],
                initialfile=f"{output_name}.docx",
                title='Save SOP Document'
            )

            root.destroy()

            if save_path:
                # Save the file
                file_stream.seek(0)
                with open(save_path, 'wb') as f:
                    f.write(file_stream.read())

                # Auto-save to archive
                import tempfile
                try:
                    # Save BPMN to temp file
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.bpmn') as tmp_bpmn:
                        tmp_bpmn.write(bpmn_content)
                        tmp_bpmn_path = tmp_bpmn.name

                    # Save DOCX to temp file
                    file_stream.seek(0)
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                        tmp_docx.write(file_stream.read())
                        tmp_docx_path = tmp_docx.name

                    # Archive the files
                    archive_manager.save_archive(
                        user_id='local',
                        process_name=output_name,
                        bpmn_file_path=tmp_bpmn_path,
                        docx_file_path=tmp_docx_path
                    )

                    # Clean up temp files
                    os.remove(tmp_bpmn_path)
                    os.remove(tmp_docx_path)
                except Exception as e:
                    debug_log(f"Error archiving: {e}")

                return jsonify({'success': True, 'message': f'Document saved to {save_path}', 'path': save_path})
            else:
                return jsonify({'success': False, 'message': 'Save cancelled'})
    return "An error occurred during file processing. Check the console for details.", 500

def start_server():
    """Start Flask server in background thread"""
    serve(app, host='127.0.0.1', port=8000, _quiet=True)

# Global reference to webview window for save dialogs
webview_window = None

def main():
    """Main entry point - launches native window with pywebview"""
    global webview_window
    import webview

    # Start Flask server in background thread
    server_thread = threading.Thread(target=start_server, daemon=True)
    server_thread.start()

    # Give server time to start
    import time
    time.sleep(0.5)

    # Create native window
    webview.create_window(
        'BPMN to SOP Generator',
        'http://127.0.0.1:8000',
        width=1000,
        height=800,
        resizable=True,
        min_size=(800, 600)
    )
    webview.start()

if __name__ == '__main__':
    main()